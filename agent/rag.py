"""
RAG 模块：加载建筑规范文档 → 递归字符切分 → 向量化存储 → 检索

支持：
1. Word (.docx) 文档加载
2. PDF (.pdf) 文本提取（pypdf）
3. 内置规则知识库 fallback（向量库不可用时自动降级）
"""
import os
from pathlib import Path
from typing import Optional

import requests
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.embeddings import Embeddings

load_dotenv()

LLM_API_KEY = os.getenv("LLM_API_KEY", "sk-your-api-key-here")
LLM_BASE_URL = os.getenv("LLM_BASE_URL", "https://api.openai.com/v1")
CHROMA_DB_DIR = Path(__file__).parent.parent / "chroma_db"
EMBEDDING_MODEL = "text-embedding-v1"


# ─── Embedding 实现 ───────────────────────────────────────────────
class _DashScopeEmbeddings(Embeddings):
    """阿里云百炼 embedding API 兼容实现（批量发送，支持进度条）"""

    def __init__(self, api_key: str, base_url: str, model: str):
        self.api_key = api_key
        self.base_url = base_url.rstrip("/")
        self.model = model

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        """批量 embedding：每次最多 25 条，带进度条"""
        url = f"{self.base_url}/embeddings"
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        all_embeddings = []
        batch_size = 25
        total = len(texts)

        # 进度条
        try:
            from tqdm import tqdm
            pbar = tqdm(total=total, desc="[RAG] Embedding", unit="chunk")
        except ImportError:
            pbar = None

        for i in range(0, total, batch_size):
            batch = texts[i:i + batch_size]
            resp = requests.post(
                url, headers=headers,
                json={"model": self.model, "input": batch},
                timeout=120,
            )
            resp.raise_for_status()
            data = resp.json()
            for item in data["data"]:
                all_embeddings.append(item["embedding"])
            if pbar:
                pbar.update(len(batch))

        if pbar:
            pbar.close()
        return all_embeddings

    def embed_query(self, text: str) -> list[float]:
        return self.embed_documents([text])[0]


_embedding = _DashScopeEmbeddings(
    api_key=LLM_API_KEY, base_url=LLM_BASE_URL, model=EMBEDDING_MODEL,
)

_text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=100,
    separators=["\n\n", "\n", "。", "；", "，", " ", ""],
)


# ─── 文档加载 ─────────────────────────────────────────────────────
def _load_docx(path: str) -> list[Document]:
    """从 .docx 文件提取文本"""
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    documents = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 5:  # 过滤太短的片段
            documents.append(Document(
                page_content=text,
                metadata={"para": i + 1, "source": path},
            ))
    return documents


def _load_pdf_text(path: str) -> list[Document]:
    """从文本型 PDF 提取文本"""
    from pypdf import PdfReader

    reader = PdfReader(path)
    documents = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            documents.append(Document(
                page_content=text.strip(),
                metadata={"page": i + 1, "source": path},
            ))
    return documents


# ─── 向量库构建 ───────────────────────────────────────────────────
def build_vectorstore(file_path: str) -> Chroma:
    """
    从文档构建向量库：加载 → 切分 → 向量化 → 存入 Chroma。

    支持 .docx 和 .pdf 格式。
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        documents = _load_docx(str(path))
        print(f"[RAG] 已加载 Word 文档：{len(documents)} 个段落")
    elif suffix == ".pdf":
        documents = _load_pdf_text(str(path))
        if not documents:
            raise RuntimeError(f"PDF 无文本内容（可能是扫描版）：{file_path}")
        print(f"[RAG] 已加载 PDF：{len(documents)} 页")
    else:
        raise ValueError(f"不支持的文件格式：{suffix}（仅支持 .docx / .pdf）")

    if not documents:
        raise RuntimeError(f"无法从文档中提取文本：{file_path}")

    chunks = _text_splitter.split_documents(documents)
    print(f"[RAG] 切分完成：{len(chunks)} 个文本块")

    CHROMA_DB_DIR.mkdir(parents=True, exist_ok=True)
    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=_embedding,
        persist_directory=str(CHROMA_DB_DIR),
    )
    print(f"[RAG] 向量库已构建：{CHROMA_DB_DIR}")
    return vectorstore


# ─── 向量库加载 ───────────────────────────────────────────────────
def load_vectorstore() -> Chroma | None:
    """
    加载已有向量库；不存在时自动从项目根目录的 docx/pdf 构建。
    构建失败返回 None。
    """
    if (CHROMA_DB_DIR / "chroma.sqlite3").exists():
        try:
            return Chroma(
                persist_directory=str(CHROMA_DB_DIR),
                embedding_function=_embedding,
            )
        except Exception as e:
            print(f"[RAG] 加载向量库失败：{e}")
            return None

    # 自动构建：优先找 .docx，其次 .pdf
    root = Path(__file__).parent.parent
    doc_files = list(root.glob("*.docx")) + list(root.glob("*.pdf"))
    if not doc_files:
        print("[RAG] 未找到文档，将使用内置规则知识库")
        return None

    target = str(doc_files[0])
    print(f"[RAG] 从 {target} 构建向量库...")
    try:
        return build_vectorstore(target)
    except Exception as e:
        print(f"[RAG] 构建失败：{e}，降级到内置规则库")
        return None


# ─── 内置规则知识库 ────────────────────────────────────────────────
_BUILTIN = {
    "混凝土": [
        "GB/T 50344-2019 第5.2节：混凝土结构检测应包括混凝土强度、碳化深度、保护层厚度、裂缝、变形等项目。",
        "GB/T 50344-2019 第5.2.3条：混凝土抗压强度检测可采用回弹法、超声回弹综合法、钻芯法等方法。",
        "GB/T 50344-2019 第5.2.6条：混凝土裂缝检测应记录裂缝的位置、形态、长度、宽度和深度。",
        "GB/T 50344-2019 第5.2.8条：混凝土构件变形检测包括挠度和倾斜。",
    ],
    "砌体": [
        "GB/T 50344-2019 第5.3节：砌体结构检测应包括砌筑块材强度、砌筑砂浆强度、砌体裂缝、变形等项目。",
        "GB/T 50344-2019 第5.3.2条：砌筑砂浆强度检测可采用回弹法、贯入法、筒压法等方法。",
        "GB/T 50344-2019 第5.3.5条：砌体裂缝检测应区分受力裂缝和非受力裂缝。",
    ],
    "钢结构": [
        "GB/T 50344-2019 第5.4节：钢结构检测应包括钢材厚度、焊缝质量、涂层厚度、变形与损伤等项目。",
        "GB/T 50344-2019 第5.4.3条：焊缝质量检测可采用超声波检测、射线检测、磁粉检测等方法。",
        "GB/T 50344-2019 第5.4.6条：钢结构防腐和防火涂层厚度应满足设计要求和相关规范规定。",
    ],
    "木结构": [
        "GB/T 50344-2019 第5.5节：木结构检测应包括木材材质、连接节点、变形与损伤、腐朽与虫蛀等项目。",
    ],
    "裂缝": [
        "GB/T 50344-2019 第6.2节：裂缝检测应记录其位置、形态、走向、长度、宽度和深度。",
        "GB/T 50344-2019 第6.2.3条：一般混凝土结构裂缝宽度限值为0.2~0.4mm。",
    ],
    "变形": [
        "GB/T 50344-2019 第6.3节：结构变形检测包括水平位移、竖向变形（挠度）、倾斜等。",
        "GB/T 50344-2019 第6.3.4条：构件挠度检测值不应超过计算跨度的1/200~1/400。",
    ],
    "加层": [
        "GB/T 50344-2019 第3.1.3条：当建筑物进行加层、改造或用途变更时，应进行结构检测鉴定。",
        "GB/T 50344-2019 第7.3节：既有建筑加层改造前，应对地基基础和上部结构进行全面检测。",
        "GB/T 50344-2019 第7.3.2条：加层改造应重点关注基础承载力、竖向构件承载力和结构整体稳定性。",
    ],
    "通用": [
        "GB/T 50344-2019 第3.1.1条：建筑结构检测应遵循先调查后检测、从整体到局部、从表观到内在的原则。",
        "GB/T 50344-2019 第3.4.1条：检测报告应包含工程概况、检测目的和依据、检测方法和仪器、检测数据和分析、检测结论和建议。",
        "GB/T 50344-2019 第7.1.1条：结构安全性评定应综合考虑承载能力、结构整体稳固性和结构耐久性。",
    ],
}

_KEYWORDS = {
    "混凝土": "混凝土", "砌体": "砌体", "砖混": "砌体", "砖": "砌体",
    "钢": "钢结构", "木": "木结构",
    "裂缝": "裂缝", "裂纹": "裂缝",
    "变形": "变形", "倾斜": "变形", "挠度": "变形",
    "加层": "加层", "扩建": "加层", "改造": "加层",
}


def _retrieve_builtin(material: str, defects: list[dict], has_extension: str, max_items: int = 5) -> str:
    matched = set()
    material_lower = material.lower() if material else ""
    for kw, cat in _KEYWORDS.items():
        if kw in material_lower:
            matched.add(cat)
    for d in (defects or []):
        dtype = d.get("type", "").lower() if isinstance(d, dict) else ""
        for kw, cat in _KEYWORDS.items():
            if kw in dtype:
                matched.add(cat)
    if has_extension and has_extension != "无加层" and has_extension != "Unknown":
        matched.add("加层")

    seen = set()
    items = []
    for cat in matched:
        for item in _BUILTIN.get(cat, []):
            if item not in seen:
                seen.add(item)
                items.append(item)
                if len(items) >= max_items:
                    break
        if len(items) >= max_items:
            break
    if not items:
        items = _BUILTIN["通用"][:max_items]

    return "\n\n".join(f"【参考规范 {i}】{item}" for i, item in enumerate(items, 1))


# ─── 检索入口 ─────────────────────────────────────────────────────
def retrieve_regulations(
    vectorstore: Optional[Chroma],
    material: str,
    defects: list[dict],
    floor: str,
    has_extension: str,
    k: int = 3,
) -> str:
    """
    检索相关规范条文。优先向量库检索，不可用时降级到内置规则库。
    """
    # 路径1：向量库检索
    if vectorstore is not None:
        try:
            if defects:
                descs = [f"{d.get('type','')}(面积{d.get('area',0):.0f}px)" for d in defects]
                defect_kw = "、".join(descs)
            else:
                defect_kw = "无明显缺陷"

            query = f"{material}结构建筑，{floor}层，隐患：{defect_kw}，检测标准与评级"
            if has_extension and has_extension != "无加层":
                query += "，存在加层情况"

            docs = vectorstore.similarity_search(query, k=k)
            if docs:
                items = []
                for i, doc in enumerate(docs, 1):
                    c = doc.page_content.strip()
                    if c:
                        items.append(f"【参考规范 {i}】{c}")
                if items:
                    return "\n\n".join(items)
        except Exception as e:
            print(f"[RAG] 向量检索失败：{e}，降级到内置规则库")

    # 路径2：内置规则库
    print("[RAG] 使用内置规则知识库")
    return _retrieve_builtin(material=material, defects=defects or [], has_extension=has_extension)
