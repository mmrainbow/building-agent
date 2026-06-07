"""历史记录路由 — /history, /history/{id}, /history/{id}/export。"""

import base64
import io

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.orm import Session

from api.auth import decode_token, get_current_user
from api.schemas import RecordResponse
from db import get_all_records, get_db, get_record_detail, get_user_by_id, get_user_records
from materials import material_to_zh, replace_material_terms

router = APIRouter(tags=["history"])


def _can_access_record(user: dict, record) -> bool:
    if user["role"] == "admin":
        return True
    return record.user_id == user["user_id"]


def _record_to_dict(record) -> dict:
    images = record.images or []
    all_defects = []
    materials = []
    floors = []
    extensions = []
    image_items = []
    for idx, img in enumerate(images):
        material_zh = material_to_zh(img.material)
        materials.append(material_zh)
        floors.append(img.floor or "")
        extensions.append(img.has_extension or "")
        defects = []
        for defect in ((img.chat_image.defects if img.chat_image else []) or []):
            item = {
                "id": len(defects) + 1,
                "type": defect.defect_type,
                "area": defect.area,
                "box": defect.box_coords,
            }
            defects.append(item)
            all_defects.append({
                "type": defect.defect_type,
                "area": defect.area,
                "box": defect.box_coords,
                "image_id": img.id,
            })
        image_items.append({
            "id": img.id,
            "name": img.image_name or f"巡检图_{idx + 1}",
            "material": material_zh,
            "floor": img.floor,
            "has_extension": img.has_extension,
            "defects": defects,
            "original_url": f"/api/history/images/{img.id}/original",
            "annotated_url": f"/api/history/images/{img.id}/annotated",
        })

    return {
        "id": record.id,
        "image_count": len(images),
        "material": ", ".join(set(m for m in materials if m)) or "未知",
        "floor": ", ".join(set(f for f in floors if f)) or "未知",
        "has_extension": ", ".join(set(e for e in extensions if e)) or "未知",
        "report": replace_material_terms(record.report),
        "defects": all_defects,
        "images": image_items,
        "created_at": record.created_at.isoformat() if record.created_at else None,
    }


def _build_export_payload(record) -> dict:
    data = _record_to_dict(record)
    lines = [
        f"# 建筑外立面巡检报告 #{record.id}",
        "",
        f"- 巡检时间：{record.created_at.isoformat() if record.created_at else ''}",
        f"- 图片数量：{data['image_count']}",
        f"- 材质：{data['material']}",
        f"- 楼层：{data['floor']}",
        f"- 加层：{data['has_extension']}",
        f"- 隐患数量：{len(data['defects'])}",
        "",
        "## 报告正文",
        "",
        record.report or "无报告",
        "",
        "## 图片检测明细",
    ]
    for img in data["images"]:
        lines.extend([
            "",
            f"### {img['name']}",
            f"- 材质：{img.get('material') or '未知'}",
            f"- 楼层：{img.get('floor') or '未知'}",
            f"- 加层：{img.get('has_extension') or '未知'}",
            f"- 隐患数：{len(img.get('defects') or [])}",
        ])
        for d in img.get("defects") or []:
            lines.append(f"  - {d.get('type')}，像素面积约 {float(d.get('area') or 0):.1f}px²")
    return {"data": data, "markdown": "\n".join(lines)}


def _safe_record_or_404(record_id: int, user: dict, db: Session):
    record = get_record_detail(db, record_id)
    if not record:
        raise HTTPException(status_code=404, detail="Record not found")
    if not _can_access_record(user, record):
        raise HTTPException(status_code=403, detail="No permission")
    return record


def _user_from_query_token(token: str | None, db: Session) -> dict:
    if not token:
        raise HTTPException(status_code=401, detail="Missing token")
    payload = decode_token(token)
    user_id = payload.get("sub")
    user = get_user_by_id(db, int(user_id)) if user_id is not None else None
    if user is None:
        raise HTTPException(status_code=401, detail="Invalid token")
    role = user.role.value if hasattr(user.role, "value") else user.role
    return {"user_id": user.id, "username": user.username, "role": role}


@router.get("/history", response_model=list[RecordResponse])
def history(
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if user["role"] == "admin":
        records = get_all_records(db, limit=limit, offset=offset)
    else:
        records = get_user_records(db, user["user_id"], limit=limit, offset=offset)
    return [_record_to_dict(record) for record in records]


@router.get("/history/{record_id}", response_model=RecordResponse)
def record_detail(
    record_id: int,
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    record = _safe_record_or_404(record_id, user, db)
    return _record_to_dict(record)


@router.get("/history/images/{image_id}/{kind}")
def history_image(
    image_id: int,
    kind: str,
    token: str | None = Query(None),
    db: Session = Depends(get_db),
):
    """获取历史巡检原图或动态生成的标注图。kind=original|annotated。"""
    import cv2
    import numpy as np
    from agent.skills.inspection_skill import draw_defects
    from db.models import ImageInspection

    user = _user_from_query_token(token, db)
    img = db.query(ImageInspection).filter(ImageInspection.id == image_id).first()
    if not img or not img.record:
        raise HTTPException(status_code=404, detail="Image not found")
    if not _can_access_record(user, img.record):
        raise HTTPException(status_code=403, detail="No permission")
    if not img.chat_image or not img.chat_image.data:
        raise HTTPException(status_code=404, detail="Image data not found")
    if kind not in {"original", "annotated"}:
        raise HTTPException(status_code=400, detail="kind must be original or annotated")

    if kind == "original":
        return Response(content=img.chat_image.data, media_type=img.chat_image.mime_type or "image/jpeg")

    nparr = np.frombuffer(img.chat_image.data, np.uint8)
    bgr = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
    if bgr is None:
        raise HTTPException(status_code=500, detail="Image decode failed")
    rgb = cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB)
    defects = [
        {"id": idx + 1, "type": d.defect_type, "area": d.area, "box": d.box_coords}
        for idx, d in enumerate(img.chat_image.defects or [])
    ]
    annotated = draw_defects(rgb, defects)
    _, buf = cv2.imencode(".jpg", cv2.cvtColor(annotated, cv2.COLOR_RGB2BGR))
    return Response(content=buf.tobytes(), media_type="image/jpeg")


@router.get("/history/{record_id}/export")
def history_export(
    record_id: int,
    format: str = Query("xlsx", pattern="^(xlsx|docx|md)$"),
    user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """导出巡检记录为 Excel / Word / Markdown 文件。"""
    record = _safe_record_or_404(record_id, user, db)
    payload = _build_export_payload(record)

    if format == "md":
        content = payload["markdown"].encode("utf-8-sig")
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename=inspection_{record_id}.md"},
        )

    if format == "docx":
        from docx import Document
        from docx.shared import Inches
        import cv2
        import numpy as np
        from agent.skills.inspection_skill import draw_defects

        doc = Document()
        doc.add_heading(f"建筑外立面巡检报告 #{record.id}", level=1)
        doc.add_paragraph(f"巡检时间：{record.created_at.isoformat() if record.created_at else ''}")
        doc.add_paragraph(f"图片数量：{payload['data']['image_count']}")
        doc.add_paragraph(f"材质：{payload['data']['material']}")
        doc.add_paragraph(f"楼层：{payload['data']['floor']}")
        doc.add_paragraph(f"加层：{payload['data']['has_extension']}")
        doc.add_heading("报告正文", level=2)
        doc.add_paragraph(record.report or "无报告")
        doc.add_heading("图片检测明细", level=2)
        for img_entry in (record.images or []):
            img_data = payload["data"]["images"]
            img_meta = next((m for m in img_data if m["id"] == img_entry.id), {})
            doc.add_heading(img_meta.get("name", f"巡检图"), level=3)
            doc.add_paragraph(f"材质：{img_meta.get('material') or '未知'}")
            doc.add_paragraph(f"楼层：{img_meta.get('floor') or '未知'}")
            doc.add_paragraph(f"加层：{img_meta.get('has_extension') or '未知'}")
            defects = img_meta.get("defects") or []
            doc.add_paragraph(f"隐患数：{len(defects)}")
            for d in defects:
                doc.add_paragraph(f"{d.get('type')}，像素面积约 {float(d.get('area') or 0):.1f}px²", style="List Bullet")
            # 嵌入标注图
            if img_entry.chat_image and img_entry.chat_image.data:
                try:
                    nparr = np.frombuffer(img_entry.chat_image.data, np.uint8)
                    bgr = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                    if bgr is not None:
                        rgb = cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB)
                        annotated = draw_defects(rgb, defects)
                        _, jpg = cv2.imencode(".jpg", cv2.cvtColor(annotated, cv2.COLOR_RGB2BGR))
                        img_stream = io.BytesIO(jpg.tobytes())
                        doc.add_picture(img_stream, width=Inches(5.5))
                        img_stream.close()
                except Exception:
                    pass
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=inspection_{record_id}.docx"},
        )

    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "巡检报告"
    ws.append(["ID", record.id])
    ws.append(["时间", record.created_at.isoformat() if record.created_at else ""])
    ws.append(["报告", record.report or ""])
    ws.append([])
    ws.append(["图片", "材质", "楼层", "加层", "隐患数"])
    for img in record.images or []:
        defects = img.chat_image.defects if img.chat_image else []
        ws.append([img.image_name, material_to_zh(img.material), img.floor, img.has_extension, len(defects or [])])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=inspection_{record_id}.xlsx"},
    )
