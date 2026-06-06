"""
RAG 模块：加载建筑规范文档 → 递归字符切分 → 向量化存储 → 检索

支持：
1. Word (.docx) 文档加载
2. PDF (.pdf) 文本提取（pypdf）
3. 内置规则知识库 fallback（向量库不可用时自动降级）
"""
import os
from pathlib import Path
from typing import Optional

import requests
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.embeddings import Embeddings

load_dotenv()

EMBEDDING_API_KEY = os.getenv("EMBEDDING_API_KEY", os.getenv("LLM_API_KEY", ""))
EMBEDDING_BASE_URL = os.getenv(
    "EMBEDDING_BASE_URL",os.getenv("LLM_BASE_URL", " "),)
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-v3")
CHROMA_DB_DIR = Path(__file__).resolve().parent.parent / "chroma_db"


# ─── Embedding 实现 ───────────────────────────────────────────────
class _DashScopeEmbeddings(Embeddings):
    """阿里云 DashScope embedding — 原生 API，不走兼容模式。"""

    def __init__(self, api_key: str, base_url: str, model: str):
        self.api_key = api_key
        self.model = model
        self.url = (
            "https://dashscope.aliyuncs.com/api/v1/services/embeddings/"
            "text-embedding/text-embedding"
        )

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        """批量 embedding：每次最多 25 条，带进度条"""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        all_embeddings = []
        batch_size = 10  # text-embedding-v3 单批上限 10 条
        total = len(texts)

        try:
            from tqdm import tqdm

            pbar = tqdm(total=total, desc="[RAG] Embedding", unit="chunk")
        except ImportError:
            pbar = None

        for i in range(0, total, batch_size):
            batch = texts[i : i + batch_size]
            resp = requests.post(
                self.url,
                headers=headers,
                json={
                    "model": self.model,
                    "input": {"texts": batch},
                    "parameters": {},
                },
                timeout=120,
            )
            resp.raise_for_status()
            data = resp.json()
            for emb in data["output"]["embeddings"]:
                all_embeddings.append(emb["embedding"])
            if pbar:
                pbar.update(len(batch))

        if pbar:
            pbar.close()
        return all_embeddings

    def embed_query(self, text: str) -> list[float]:
        return self.embed_documents([text])[0]


_embedding = _DashScopeEmbeddings(
    api_key=EMBEDDING_API_KEY, base_url=EMBEDDING_BASE_URL, model=EMBEDDING_MODEL,
)

_text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=100,
    separators=["\n\n", "\n", "。", "；", "，", " ", ""],
)


# ─── 文档加载 ─────────────────────────────────────────────────────
def _load_docx(path: str) -> list[Document]:
    """从 .docx 文件提取文本"""
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    documents = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 5:  # 过滤太短的片段
            documents.append(Document(
                page_content=text,
                metadata={"para": i + 1, "source": path},
            ))
    return documents


def _load_pdf_text(path: str) -> list[Document]:
    """从文本型 PDF 提取文本"""
    from pypdf import PdfReader

    reader = PdfReader(path)
    documents = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            documents.append(Document(
                page_content=text.strip(),
                metadata={"page": i + 1, "source": path},
            ))
    return documents


# ─── 向量库构建 ───────────────────────────────────────────────────
def build_vectorstore(file_path: str) -> Chroma:
    """
    从文档构建向量库：加载 → 切分 → 向量化 → 存入 Chroma。

    支持 .docx 和 .pdf 格式。
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        documents = _load_docx(str(path))
        print(f"[RAG] 已加载 Word 文档：{len(documents)} 个段落")
    elif suffix == ".pdf":
        documents = _load_pdf_text(str(path))
        if not documents:
            raise RuntimeError(f"PDF 无文本内容（可能是扫描版）：{file_path}")
        print(f"[RAG] 已加载 PDF：{len(documents)} 页")
    else:
        raise ValueError(f"不支持的文件格式：{suffix}（仅支持 .docx / .pdf）")

    if not documents:
        raise RuntimeError(f"无法从文档中提取文本：{file_path}")

    chunks = _text_splitter.split_documents(documents)
    print(f"[RAG] 切分完成：{len(chunks)} 个文本块")

    CHROMA_DB_DIR.mkdir(parents=True, exist_ok=True)
    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=_embedding,
        persist_directory=str(CHROMA_DB_DIR),
    )
    print(f"[RAG] 向量库已构建：{CHROMA_DB_DIR}")
    return vectorstore


# ─── 向量库加载 ───────────────────────────────────────────────────
def load_vectorstore() -> Chroma | None:
    """
    加载已有向量库；不存在时自动从项目根目录的 docx/pdf 构建。
    构建失败返回 None。
    """
    if (CHROMA_DB_DIR / "chroma.sqlite3").exists():
        try:
            return Chroma(
                persist_directory=str(CHROMA_DB_DIR),
                embedding_function=_embedding,
            )
        except Exception as e:
            print(f"[RAG] 加载向量库失败：{e}")
            return None

    # 自动构建：从 rag_data/ 目录优先找 .docx，其次 .pdf
    root = Path(__file__).parent.parent.parent
    rag_dir = root / "rag_data"
    doc_files = list(rag_dir.glob("*.docx")) + list(rag_dir.glob("*.pdf"))
    if not doc_files:
        print("[RAG] 未找到文档，将使用内置规则知识库")
        return None

    target = str(doc_files[0])
    print(f"[RAG] 从 {target} 构建向量库...")
    try:
        return build_vectorstore(target)
    except Exception as e:
        print(f"[RAG] 构建失败：{e}，降级到内置规则库")
        return None


# ─── 检索入口 ─────────────────────────────────────────────────────

def search_regulations(query: str, k: int = 5) -> str:
    """语义检索 — 接受自由文本查询，从 ChromaDB 检索相关规范条文。

    ReAct Agent 的 search_knowledge Tool 调用此函数。
    向量库不可用时返回空字符串。
    """
    vs = load_vectorstore()
    if vs is None:
        return ""

    try:
        docs = vs.similarity_search(query, k=k)
        if not docs:
            return ""

        items = []
        for i, doc in enumerate(docs, 1):
            c = doc.page_content.strip()
            if not c:
                continue
            meta = doc.metadata
            source = ""
            if meta.get("para"):
                source = f" [第{meta['para']}段]"
            elif meta.get("page"):
                source = f" [第{meta['page']}页]"
            items.append(f"【规范 {i}{source}】{c}")

        if items:
            print(f"[RAG] search_regulations('{query[:40]}...') → {len(items)} 条")
            return "\n\n".join(items)
        return ""
    except Exception as e:
        print(f"[RAG] search_regulations 失败: {e}")
        return ""


def retrieve_regulations(
    vectorstore: Optional[Chroma],
    material: str,
    defects: list[dict],
    floor: str,
    has_extension: str,
    k: int = 5,
) -> str:
    """多维度检索 + 去重 + 重排序，返回更丰富的规范条文。"""
    if vectorstore is None:
        print("[RAG] 向量库不可用")
        return ""

    try:
        # 1. 构建多维度查询
        queries = _build_queries(material, defects, floor, has_extension)

        # 2. 多路检索 + 去重
        seen = set()
        all_docs = []
        for q in queries:
            docs = vectorstore.similarity_search(q, k=k)
            for doc in docs:
                key = doc.page_content.strip()[:80]  # 用前 80 字做去重标识
                if key not in seen:
                    seen.add(key)
                    all_docs.append(doc)

        if not all_docs:
            print("[RAG] 未检索到相关规范")
            return ""

        # 3. 按检索先后排列（多查询自然形成优先级），保留最多 2k 条
        all_docs = all_docs[:2 * k]

        # 4. 格式化输出：带元数据
        items = []
        for i, doc in enumerate(all_docs, 1):
            c = doc.page_content.strip()
            if not c:
                continue
            meta = doc.metadata
            source_info = ""
            if meta.get("para"):
                source_info = f" [第{meta['para']}段]"
            elif meta.get("page"):
                source_info = f" [第{meta['page']}页]"
            items.append(f"【参考规范 {i}{source_info}】{c}")

        print(f"[RAG] 检索到 {len(items)} 条规范条文（去重后）")
        return "\n\n".join(items)

    except Exception as e:
        print(f"[RAG] 向量检索失败：{e}")
        return ""


def _build_queries(material: str, defects: list[dict], floor: str, has_extension: str) -> list[str]:
    """根据检测结果构建多个维度的检索查询，多角度覆盖规范条文。"""
    queries = []

    # 1. 材料 + 结构类型
    queries.append(f"{material}结构建筑 检测标准 技术规范 评级评定")

    # 2. 按缺陷类型分别检索
    if defects:
        defect_types = list(set(d.get("type", "") for d in defects if d.get("type")))
        for dt in defect_types:
            queries.append(f"建筑结构 {dt} 缺陷检测 评定标准 规范要求")
        # 综合缺陷检索
        descs = [f"{d.get('type','')}" for d in defects]
        queries.append(f"建筑结构 {'、'.join(descs[:3])} 缺陷等级 安全评估 规范")
    else:
        queries.append("建筑结构 完好 无明显缺陷 检测评定")

    # 3. 楼层相关
    floor_num = _parse_floor_num(floor)
    if floor_num:
        if floor_num <= 3:
            queries.append(f"低层建筑 多层建筑 {floor_num}层 结构检测 规范要求")
        elif floor_num <= 7:
            queries.append(f"多层建筑 {floor_num}层 结构检测 技术标准")
        elif floor_num <= 30:
            queries.append(f"高层建筑 {floor_num}层 结构安全 检测规范")
        else:
            queries.append(f"超高层建筑 {floor_num}层 结构检测 技术标准")

    # 4. 加层/改建
    if has_extension and has_extension not in ("无加层", "Unknown", "未知"):
        queries.append("建筑加层 改建 结构安全性检测 规范标准")
        queries.append("既有建筑 加层改造 承载能力评定 检测要求")

    # 5. 通用检测标准
    queries.append("建筑结构检测 评定等级 安全分类 规范条文")

    return queries


def _parse_floor_num(floor: str) -> int:
    """从楼层字符串中提取数字。"""
    import re
    match = re.search(r'(\d+)', str(floor))
    return int(match.group(1)) if match else 0


# ─── Memory 向量存储 ──────────────────────────────────────────────

_MEMORY_COLLECTION = "building_memories"


def _get_memory_vs() -> Chroma | None:
    """获取 Memory 专用的 ChromaDB 集合。不可用时返回 None。"""
    try:
        return Chroma(
            collection_name=_MEMORY_COLLECTION,
            persist_directory=str(CHROMA_DB_DIR),
            embedding_function=_embedding,
        )
    except Exception as e:
        print(f"[Memory-Vector] ChromaDB 连接失败: {e}")
        return None


def index_memory_vector(memory_id: int, content: str, user_id: int, conversation_id: int) -> str | None:
    """将记忆 Embedding 后存入 ChromaDB，返回 chroma_id。失败返回 None。"""
    vs = _get_memory_vs()
    if vs is None:
        return None
    try:
        ids = [f"mem_{memory_id}"]
        vs.add_texts(
            texts=[content],
            metadatas=[{"user_id": user_id, "conversation_id": conversation_id, "memory_id": memory_id}],
            ids=ids,
        )
        return ids[0]
    except Exception as e:
        print(f"[Memory-Vector] 索引失败 memory_id={memory_id}: {e}")
        return None


def search_memories_semantic(
    query: str, user_id: int, conversation_id: int, k: int = 10
) -> list[dict]:
    """语义检索记忆 → [{memory_id, score, content}], 按相似度排序。"""
    vs = _get_memory_vs()
    if vs is None:
        return []
    try:
        results = vs.similarity_search_with_score(
            query, k=k * 2,
            filter={"conversation_id": conversation_id},
        )
        seen = set()
        items = []
        for doc, score in results:
            mid = doc.metadata.get("memory_id")
            if mid is None or mid in seen:
                continue
            seen.add(mid)
            # ChromaDB 返回 L2 distance, 转换为 0-1 相似度
            relevance = 1.0 / (1.0 + score)
            items.append({
                "memory_id": mid,
                "relevance": round(relevance, 4),
                "content": doc.page_content[:200],
            })
        return items[:k]
    except Exception as e:
        print(f"[Memory-Vector] 检索失败: {e}")
        return []


def delete_memory_vector(memory_id: int) -> None:
    """从 ChromaDB 删除一条记忆向量。"""
    vs = _get_memory_vs()
    if vs is None:
        return
    try:
        vs.delete(ids=[f"mem_{memory_id}"])
    except Exception:
        pass


def migrate_memories_to_chroma(db) -> int:
    """将 chroma_id=NULL 的记忆逐条回填向量索引。返回迁移条数。"""
    from db.models import ConversationMemory
    orphans = db.query(ConversationMemory).filter(ConversationMemory.chroma_id.is_(None)).all()
    if not orphans:
        return 0
    print(f"[Memory-Vector] 开始迁移 {len(orphans)} 条记忆 ...")
    count = 0
    for mem in orphans:
        cid = index_memory_vector(mem.id, mem.content, mem.user_id, mem.conversation_id or 0)
        if cid:
            mem.chroma_id = cid
            count += 1
    db.commit()
    print(f"[Memory-Vector] 迁移完成: {count}/{len(orphans)} 条")
    return count
